"""Tarea 11 — Objetivos tácticos e informes (GameMetrics). Solo evidencia del repo."""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent
DOCX = OUT / "Tarea_11_Objetivos_Tacticos_GameMetrics.docx"
PDF = OUT / "Tarea_11_Objetivos_Tacticos_GameMetrics.pdf"
FONT = "Arial Narrow"


def style(doc: Document) -> None:
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    n = doc.styles["Normal"]
    n.font.name = FONT
    n._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing = 1.15
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def run(p, text: str, *, bold=False, size=Pt(11)):
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r.font.size = size
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def title(doc, text: str, size=Pt(14)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, text, bold=True, size=size)


def h(doc, text: str, size=Pt(12)):
    p = doc.add_paragraph()
    run(p, text, bold=True, size=size)


def body(doc, text: str):
    p = doc.add_paragraph()
    run(p, text)


def ficha(doc, rows: list[tuple[str, str]]):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (a, b) in enumerate(rows):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        for cell in t.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(11)
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.add_paragraph()


# (departamento, objetivo, informe, es_simple)
# Un mismo objetivo se repite cuando genera varios informes (1:N).
# Evidencia: rutas/endpoints/tablas reales del repositorio GameMetrics.
ROWS = [
    # --- Administración: 1 objetivo → 1 simple ---
    (
        "Administración de Plataforma",
        "Priorizar las solicitudes de propiedad de juegos que todavía no han sido revisadas.",
        "Listado de solicitudes de propiedad pendientes (juego, estudio, correo y fecha), ordenadas desde la más reciente.",
        True,
    ),
    # --- Administración: 1 objetivo → 2 simples ---
    (
        "Administración de Plataforma",
        "Controlar el personal y los acuerdos comerciales registrados por la empresa.",
        "Listado de empleados activos con nombre, cargo, departamento y fecha de ingreso.",
        True,
    ),
    (
        "Administración de Plataforma",
        "Controlar el personal y los acuerdos comerciales registrados por la empresa.",
        "Listado de contratos con publicador, tipo, fechas, valor y estado.",
        True,
    ),
    # --- Administración: 1 objetivo → 1 simple + 3 compuestos ---
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Listado de liquidaciones ya pagadas a estudios (partner, monto, método y referencia).",
        True,
    ),
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Valor bruto total de ventas acumulado (GMV) en dólares.",
        False,
    ),
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Ingresos acumulados de GameMetrics por comisión de plataforma.",
        False,
    ),
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Monto total adeudado a estudios publicadores pendiente de liquidar.",
        False,
    ),
    # --- Administración: ETL ---
    (
        "Administración de Plataforma",
        "Verificar si los trabajos de carga de datos están listos antes de actualizar el catálogo analítico.",
        "Estado actual de los trabajos de carga (dataset, dimensiones, tablas en tiempo real y catálogo).",
        True,
    ),
    # --- Ventas: 1 objetivo → 2 simples ---
    (
        "Ventas y Marketing",
        "Organizar las campañas promocionales y los títulos autorizados para venta.",
        "Listado de campañas de marketing con nombre, juego, presupuesto, canal y estado.",
        True,
    ),
    (
        "Ventas y Marketing",
        "Organizar las campañas promocionales y los títulos autorizados para venta.",
        "Listado del catálogo de distribución con juego, plataforma, precio, región y estado.",
        True,
    ),
    # --- Ventas: featured + desempeño ---
    (
        "Ventas y Marketing",
        "Controlar las promociones de visibilidad pagadas en la tienda.",
        "Listado de ubicaciones destacadas activas por juego y estudio.",
        True,
    ),
    (
        "Ventas y Marketing",
        "Comparar el desempeño comercial entre estudios y juegos para detectar los que más venden.",
        "Unidades vendidas, ingreso bruto, comisión e ingreso neto por cada estudio publicador.",
        False,
    ),
    (
        "Ventas y Marketing",
        "Comparar el desempeño comercial entre estudios y juegos para detectar los que más venden.",
        "Cantidad de reembolsos asociados a cada estudio publicador.",
        False,
    ),
    # --- Analytics: 1 objetivo → varios compuestos ---
    (
        "Analytics BI",
        "Verificar el volumen y la composición del catálogo analítico cargado por semana.",
        "Cantidad total de videojuegos disponibles en la semana seleccionada.",
        False,
    ),
    (
        "Analytics BI",
        "Verificar el volumen y la composición del catálogo analítico cargado por semana.",
        "Cantidad de juegos agrupados por género en la semana seleccionada.",
        False,
    ),
    (
        "Analytics BI",
        "Verificar el volumen y la composición del catálogo analítico cargado por semana.",
        "Cantidad de juegos agrupados por plataforma en la semana seleccionada.",
        False,
    ),
    (
        "Analytics BI",
        "Detectar los títulos mejor valorados para decidir destacados editoriales en la tienda.",
        "Listado de los 10 juegos con mayor calificación en la semana seleccionada.",
        False,
    ),
    # --- Comercio: 1 objetivo → 2 simples ---
    (
        "Comercio digital",
        "Permitir al jugador consultar y gestionar las compras que ya realizó.",
        "Listado de juegos comprados en la biblioteca del usuario (nombre, precio y fecha).",
        True,
    ),
    (
        "Comercio digital",
        "Permitir al jugador consultar y gestionar las compras que ya realizó.",
        "Listado de compras elegibles para devolución según la política de 14 días.",
        True,
    ),
    # --- Atención: 1 simple (solo lo implementado) ---
    (
        "Atención al Cliente",
        "Atender primero las solicitudes de ayuda que siguen abiertas.",
        "Listado de tickets de soporte del usuario con asunto, prioridad, estado y fecha.",
        True,
    ),
    # --- Publishers: 1 simple ---
    (
        "Distribución B2B (estudios publicadores)",
        "Verificar el estado de los juegos que el estudio ha solicitado administrar.",
        "Listado de juegos reclamados por el estudio con estado pendiente, aprobado o rechazado.",
        True,
    ),
    # --- Publishers: 1 objetivo → 1 simple + 2 compuestos ---
    (
        "Distribución B2B (estudios publicadores)",
        "Controlar cuánto ha ganado el estudio, cuánto puede cobrar y qué ya le pagaron.",
        "Listado de liquidaciones recibidas por el estudio (monto, fecha y referencia).",
        True,
    ),
    (
        "Distribución B2B (estudios publicadores)",
        "Controlar cuánto ha ganado el estudio, cuánto puede cobrar y qué ya le pagaron.",
        "Resumen de ingresos del estudio: bruto, comisión de plataforma, neto y saldo disponible.",
        False,
    ),
    (
        "Distribución B2B (estudios publicadores)",
        "Controlar cuánto ha ganado el estudio, cuánto puede cobrar y qué ya le pagaron.",
        "Ingresos netos y ventas agrupados por cada juego del estudio.",
        False,
    ),
]

SIMPLE_FICHAS = [
    (
        "Administración de Plataforma",
        "Priorizar las solicitudes de propiedad de juegos que todavía no han sido revisadas.",
        "Listado de solicitudes de propiedad pendientes (juego, estudio, correo y fecha), ordenadas desde la más reciente.",
        "¿Qué solicitudes de propiedad de juegos siguen sin decidir?",
        "Fila por solicitud: nombre del juego, compañía, correo, estado pendiente y fecha.",
        "Consulta fact_partner_games con submission_status = pending; enriquece con fact_partner_accounts.",
        "Vista Admin en /admin · API GET /admin/game-claims",
        "Aprobar o rechazar la solicitud.",
    ),
    (
        "Administración de Plataforma",
        "Controlar el personal y los acuerdos comerciales registrados por la empresa.",
        "Listado de empleados activos con nombre, cargo, departamento y fecha de ingreso.",
        "¿Quiénes están registrados como personal de la empresa?",
        "Lista paginada de empleados no eliminados.",
        "Consulta emp_records con collection = empleados y deleted = false.",
        "Vista Empresa en /empresa · colección empleados",
        "Actualizar cargos o registrar altas y bajas.",
    ),
    (
        "Administración de Plataforma",
        "Controlar el personal y los acuerdos comerciales registrados por la empresa.",
        "Listado de contratos con publicador, tipo, fechas, valor y estado.",
        "¿Qué acuerdos comerciales existen y en qué estado están?",
        "Lista paginada de contratos con publicador, fechas y estado.",
        "Consulta emp_records con collection = contratos y deleted = false.",
        "Vista Empresa en /empresa · colección contratos",
        "Renovar, modificar o cerrar un acuerdo.",
    ),
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Listado de liquidaciones ya pagadas a estudios (partner, monto, método y referencia).",
        "¿Qué pagos a estudios ya quedaron registrados?",
        "Historial de payouts con monto, método y referencia.",
        "Consulta fact_partner_payouts ordenada por fecha.",
        "Vista Admin en /admin · API GET /admin/payouts",
        "Auditar referencias y conciliar pagos.",
    ),
    (
        "Administración de Plataforma",
        "Verificar si los trabajos de carga de datos están listos antes de actualizar el catálogo analítico.",
        "Estado actual de los trabajos de carga (dataset, dimensiones, tablas en tiempo real y catálogo).",
        "¿Los trabajos de carga están idle, running, ok o en error?",
        "Estado por cada trabajo del panel ETL.",
        "Lectura del servidor ETL (etl_server.py) y semanas cargadas.",
        "Dashboard en / · pestaña ETL",
        "Ejecutar o reintentar una carga antes de usar informes analíticos.",
    ),
    (
        "Ventas y Marketing",
        "Organizar las campañas promocionales y los títulos autorizados para venta.",
        "Listado de campañas de marketing con nombre, juego, presupuesto, canal y estado.",
        "¿Qué campañas existen y cuál es su estado?",
        "Lista de campañas con presupuesto, canal y fechas.",
        "Consulta emp_records con collection = campanas_marketing.",
        "Vista Empresa en /empresa · colección campanas_marketing",
        "Activar, pausar o planificar una campaña.",
    ),
    (
        "Ventas y Marketing",
        "Organizar las campañas promocionales y los títulos autorizados para venta.",
        "Listado del catálogo de distribución con juego, plataforma, precio, región y estado.",
        "¿Qué títulos están autorizados para comercializarse?",
        "Lista de juegos del catálogo comercial interno.",
        "Consulta emp_records con collection = catalogo_distribucion.",
        "Vista Empresa en /empresa · colección catalogo_distribucion",
        "Autorizar o retirar un título de la oferta.",
    ),
    (
        "Ventas y Marketing",
        "Controlar las promociones de visibilidad pagadas en la tienda.",
        "Listado de ubicaciones destacadas activas por juego y estudio.",
        "¿Qué juegos tienen promoción de visibilidad activa?",
        "Placements activos por partner y producto.",
        "Consulta fact_featured_placements.",
        "Panel Partner en /my-partner · tienda /store/featured",
        "Renovar o cancelar una promoción pagada.",
    ),
    (
        "Comercio digital",
        "Permitir al jugador consultar y gestionar las compras que ya realizó.",
        "Listado de juegos comprados en la biblioteca del usuario (nombre, precio y fecha).",
        "¿Qué juegos compró este usuario?",
        "Compras no reembolsadas del usuario autenticado.",
        "Consulta fact_purchases filtrando user_id y refunded = false.",
        "Vista Biblioteca en /my-library",
        "Instalar, jugar o iniciar una devolución.",
    ),
    (
        "Comercio digital",
        "Permitir al jugador consultar y gestionar las compras que ya realizó.",
        "Listado de compras elegibles para devolución según la política de 14 días.",
        "¿Cuáles compras todavía pueden devolverse?",
        "Compras dentro del plazo y sin reembolso previo.",
        "Consulta fact_purchases con reglas del módulo refunds.",
        "Wizard de reembolso en /my-library",
        "Procesar o rechazar la devolución.",
    ),
    (
        "Atención al Cliente",
        "Atender primero las solicitudes de ayuda que siguen abiertas.",
        "Listado de tickets de soporte del usuario con asunto, prioridad, estado y fecha.",
        "¿Qué solicitudes de ayuda tiene abiertas el usuario?",
        "Tickets del usuario ordenados por fecha.",
        "Consulta fact_support_tickets filtrando user_id.",
        "Vista Soporte en /my-support · API GET /support/tickets",
        "Responder o cerrar el ticket.",
    ),
    (
        "Distribución B2B (estudios publicadores)",
        "Verificar el estado de los juegos que el estudio ha solicitado administrar.",
        "Listado de juegos reclamados por el estudio con estado pendiente, aprobado o rechazado.",
        "¿Qué juegos reclamó el estudio y en qué estado están?",
        "Claims del partner con submission_status.",
        "Consulta fact_partner_games filtrando partner_id.",
        "Panel Partner en /my-partner",
        "Esperar aprobación o reclamar otro título.",
    ),
    (
        "Distribución B2B (estudios publicadores)",
        "Controlar cuánto ha ganado el estudio, cuánto puede cobrar y qué ya le pagaron.",
        "Listado de liquidaciones recibidas por el estudio (monto, fecha y referencia).",
        "¿Qué pagos ya recibió este estudio?",
        "Historial de payouts del partner autenticado.",
        "Consulta fact_partner_payouts filtrando partner_id.",
        "Sección Payouts en /my-partner",
        "Conciliar pagos con el saldo disponible.",
    ),
]

COMPOUND_FICHAS = [
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Valor bruto total de ventas acumulado (GMV) en dólares.",
        "¿Cuánto dinero en ventas brutas ha movido la plataforma?",
        "Suma de montos brutos de ventas en el libro de ingresos (menos efecto de reembolsos).",
        "fact_partner_ledger (gross_amount, entry_type).",
        "KPI GMV en /admin · API GET /admin/dashboard",
        "Evaluar crecimiento comercial de la plataforma.",
    ),
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Ingresos acumulados de GameMetrics por comisión de plataforma.",
        "¿Cuánto retuvo GameMetrics por su comisión?",
        "Suma de platform_fee_amount en el libro de ingresos.",
        "fact_partner_ledger (platform_fee_amount).",
        "KPI Ingresos GameMetrics en /admin",
        "Proyectar ingresos propios de la empresa.",
    ),
    (
        "Administración de Plataforma",
        "Medir el dinero que mueve la plataforma, lo que retiene GameMetrics y lo que aún se debe a los estudios.",
        "Monto total adeudado a estudios publicadores pendiente de liquidar.",
        "¿Cuánto dinero todavía se debe pagar a los estudios?",
        "Agregación de neto publisher menos liquidaciones ya pagadas (con hold y mínimo).",
        "fact_partner_ledger, fact_partner_payouts.",
        "KPI Adeudado publishers en /admin",
        "Programar la siguiente liquidación.",
    ),
    (
        "Ventas y Marketing",
        "Comparar el desempeño comercial entre estudios y juegos para detectar los que más venden.",
        "Unidades vendidas, ingreso bruto, comisión e ingreso neto por cada estudio publicador.",
        "¿Qué estudios venden más y cuánto generan?",
        "Agrupación por partner_id sumando unidades, bruto, fee y neto.",
        "fact_partner_ledger, fact_partner_accounts.",
        "Tabla de partners en /admin/dashboard",
        "Priorizar soporte comercial a los estudios con más ventas.",
    ),
    (
        "Ventas y Marketing",
        "Comparar el desempeño comercial entre estudios y juegos para detectar los que más venden.",
        "Cantidad de reembolsos asociados a cada estudio publicador.",
        "¿Qué estudios concentran más devoluciones?",
        "Conteo de asientos de tipo refund por partner_id.",
        "fact_partner_ledger (entry_type = refund).",
        "Columna refund_count en /admin/dashboard",
        "Revisar títulos con alta tasa de devolución.",
    ),
    (
        "Analytics BI",
        "Verificar el volumen y la composición del catálogo analítico cargado por semana.",
        "Cantidad total de videojuegos disponibles en la semana seleccionada.",
        "¿Cuántos juegos hay cargados hasta la semana N?",
        "COUNT(*) sobre fact_videogames filtrado por semana.",
        "fact_videogames (semana) cargado por ETL.",
        "Dashboard / · GET /games/count?semana=N",
        "Confirmar que la carga semanal terminó.",
    ),
    (
        "Analytics BI",
        "Verificar el volumen y la composición del catálogo analítico cargado por semana.",
        "Cantidad de juegos agrupados por género en la semana seleccionada.",
        "¿Qué géneros concentran más títulos?",
        "Agrupación y conteo por género sobre el catálogo OFFLINE.",
        "fact_videogames, dim_generos.",
        "Pestaña Analytics · GET /dashboard/by-genre",
        "Orientar decisiones de catálogo por género dominante.",
    ),
    (
        "Analytics BI",
        "Verificar el volumen y la composición del catálogo analítico cargado por semana.",
        "Cantidad de juegos agrupados por plataforma en la semana seleccionada.",
        "¿Qué plataformas concentran más títulos?",
        "Agrupación y conteo por plataforma.",
        "fact_videogames, dim_plataformas.",
        "Pestaña Analytics · GET /dashboard/by-platform",
        "Priorizar plataformas con mayor cobertura.",
    ),
    (
        "Analytics BI",
        "Detectar los títulos mejor valorados para decidir destacados editoriales en la tienda.",
        "Listado de los 10 juegos con mayor calificación en la semana seleccionada.",
        "¿Cuáles son los 10 juegos mejor valorados?",
        "Ordenar por rating y limitar a 10, filtrando por semana.",
        "fact_videogames (name, rating, metacritic).",
        "Analytics · GET /dashboard/top-rated",
        "Elegir títulos para el carrusel editorial.",
    ),
    (
        "Distribución B2B (estudios publicadores)",
        "Controlar cuánto ha ganado el estudio, cuánto puede cobrar y qué ya le pagaron.",
        "Resumen de ingresos del estudio: bruto, comisión de plataforma, neto y saldo disponible.",
        "¿Cuánto ha ganado este estudio y cuánto puede retirar?",
        "Agregación del libro de ingresos del partner (gross, fee, net, available).",
        "fact_partner_ledger, fact_partner_payouts, fact_partner_accounts.",
        "Panel /my-partner · resumen de earnings",
        "Solicitar liquidación cuando el saldo supera el mínimo.",
    ),
    (
        "Distribución B2B (estudios publicadores)",
        "Controlar cuánto ha ganado el estudio, cuánto puede cobrar y qué ya le pagaron.",
        "Ingresos netos y ventas agrupados por cada juego del estudio.",
        "¿Qué juego del estudio genera más ingresos?",
        "Agrupación por product_id dentro del ledger del partner.",
        "fact_partner_ledger, fact_partner_games.",
        "Sección de juegos/revenue en /my-partner",
        "Invertir más esfuerzo en el juego de mejor rendimiento.",
    ),
]


def build() -> Document:
    doc = Document()
    style(doc)

    # 1. Portada
    title(doc, "[Institución — editable]")
    title(doc, "[Carrera — editable]")
    doc.add_paragraph()
    title(doc, "Construcción del Software")
    title(doc, "Tarea 11")
    doc.add_paragraph()
    title(doc, "Análisis de objetivos tácticos e informes")
    title(doc, "Proyecto: GameMetrics S.A.")
    doc.add_paragraph()
    body(doc, "Integrantes: [editable]")
    body(doc, "Docente: [editable]")
    body(doc, "Semestre: [editable]")
    body(doc, "Fecha: [editable]")
    doc.add_page_break()

    # 2. Operativa → táctica (corto)
    h(doc, "1. De la fase operativa a la fase táctica", Pt(14))
    body(
        doc,
        "GameMetrics S.A. ya opera registro, tienda, compras, biblioteca, estudios publicadores, "
        "panel admin y carga ETL. En la fase táctica, cada área necesita informes para dirigir "
        "y controlar esos procesos. Para efectos del análisis táctico, los procesos se agrupan "
        "en áreas funcionales; no representan un organigrama legal.",
    )
    body(
        doc,
        "El proyecto no usa PostgreSQL. La capa operacional (equivalente a base relacional) es "
        "Pinot REALTIME vía Kafka: estado actual de empleados, contratos, tickets, compras y "
        "partners. La capa analítica (base columnar) es Pinot OFFLINE más agregaciones OLAP, "
        "alimentada por el pipeline ETL.",
    )

    # 3. Relación
    h(doc, "2. Relación entre áreas, objetivos e informes", Pt(14))
    body(
        doc,
        "Área funcional 1:N objetivos tácticos. Objetivo táctico 1:N informes. "
        "Un mismo objetivo puede generar un informe simple, varios simples, un compuesto, "
        "varios compuestos, o simples y compuestos a la vez. Cada informe aporta una "
        "perspectiva distinta. En la tabla, el objetivo se repite en filas separadas.",
    )
    body(
        doc,
        "Ejemplo en este documento: el objetivo de medir el dinero de la plataforma genera "
        "cuatro informes (listado de liquidaciones + GMV + comisión GameMetrics + adeudado). "
        "El objetivo del estudio publicador sobre sus ganancias genera tres (historial de "
        "pagos + resumen de ingresos + ingresos por juego).",
    )

    # 4. Criterio
    h(doc, "3. Criterio de clasificación", Pt(14))
    body(
        doc,
        "Informe simple: listado o consulta directa sobre el estado actual (filtro, "
        "ordenamiento). No requiere agregación analítica previa. En GameMetrics se sirve "
        "desde Pinot REALTIME / emp_records y ya está visible en una lista o panel.",
    )
    body(
        doc,
        "Informe compuesto: suma, conteo, agrupación o indicador calculado. Se consulta "
        "desde la base columnar (Pinot OFFLINE o agregados del libro de ingresos B2B).",
    )
    body(
        doc,
        "Solo se incluyen objetivos e informes respaldados por tablas, endpoints y vistas "
        "ya existentes en el repositorio. No se inventan indicadores sin datos operativos.",
    )

    # 5. Tabla principal (entregable central)
    h(doc, "4. Tabla principal (entregable central)", Pt(14))
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, txt in enumerate(
        ["DEPARTAMENTO", "OBJETIVOS TÁCTICOS", "¿ES UN INFORME SIMPLE?", "¿ES UN INFORME COMPUESTO?"]
    ):
        table.rows[0].cells[i].text = txt
    for dep, obj, inf, simple in ROWS:
        cells = table.add_row().cells
        cells[0].text = dep
        cells[1].text = f"Objetivo:\n{obj}\n\nInforme:\n{inf}"
        cells[2].text = "X" if simple else ""
        cells[3].text = "" if simple else "X"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.add_paragraph()
    body(
        doc,
        f"Total filas (informes): {len(ROWS)}. "
        f"Simples: {sum(1 for *_, s in ROWS if s)}. "
        f"Compuestos: {sum(1 for *_, s in ROWS if not s)}. "
        f"Áreas: {len({r[0] for r in ROWS})}. "
        f"Objetivos distintos: {len({r[1] for r in ROWS})}.",
    )

    # 6. Informes simples
    h(doc, "5. Informes simples (base operacional / listados implementados)", Pt(14))
    body(
        doc,
        "Todos los informes simples de la tabla están implementados como listas o paneles "
        "en el sistema (requisito del docente).",
    )
    cur_a = cur_o = ""
    for area, obj, name, q, resp, how, where, dec in SIMPLE_FICHAS:
        if area != cur_a:
            cur_a = area
            cur_o = ""
            h(doc, area)
        if obj != cur_o:
            cur_o = obj
            body(doc, f"Objetivo: {obj}")
        ficha(
            doc,
            [
                ("INFORME SIMPLE", name),
                ("PREGUNTA DE NEGOCIO", q),
                ("RESPUESTA QUE RECIBE EL RESPONSABLE", resp),
                ("CÓMO LO OBTIENE EL SISTEMA", how),
                ("EVIDENCIA EN EL SISTEMA", where),
                ("DECISIÓN QUE PERMITE", dec),
            ],
        )

    # 7. Informes compuestos
    h(doc, "6. Informes compuestos (base columnar / agregaciones)", Pt(14))
    cur_a = cur_o = ""
    for area, obj, name, q, resp, how, where, dec in COMPOUND_FICHAS:
        if area != cur_a:
            cur_a = area
            cur_o = ""
            h(doc, area)
        if obj != cur_o:
            cur_o = obj
            body(doc, f"Objetivo: {obj}")
        ficha(
            doc,
            [
                ("INFORME COMPUESTO", name),
                ("PREGUNTA DE NEGOCIO", q),
                ("RESPUESTA QUE RECIBE EL RESPONSABLE", resp),
                ("CÓMO LO CALCULA EL SISTEMA", how),
                ("EVIDENCIA EN EL SISTEMA", where),
                ("DECISIÓN QUE PERMITE", dec),
            ],
        )

    # 8. BDR / BDC
    h(doc, "7. Diferencia entre capa operacional y capa columnar", Pt(14))
    body(
        doc,
        "Capa operacional (rol de base relacional en GameMetrics — Pinot REALTIME): mantiene "
        "operaciones diarias, registra el estado actual, permite crear/actualizar/consultar "
        "y responde informes simples.",
    )
    body(
        doc,
        "Capa columnar (Pinot OFFLINE + agregados OLAP): conserva datos preparados para "
        "análisis, facilita sumas, conteos y agrupaciones, y responde informes compuestos. "
        "No sustituye a la operacional; ambas son complementarias.",
    )

    # 9. ETL / Airflow
    h(doc, "8. ETL y Airflow", Pt(14))
    body(
        doc,
        "Flujo de informes compuestos: fuentes (RAWG / Parquet / operaciones) → extracción → "
        "validación → limpieza → transformación → carga en Pinot OFFLINE → publicación en "
        "dashboard. Los scripts viven en etl/ (00–09 y siguientes) y se disparan desde "
        "etl_server.py / panel ETL.",
    )
    body(
        doc,
        "[Operación REALTIME] → [ETL Python] → [Pinot OFFLINE / agregados] → [Informe táctico]",
    )
    body(
        doc,
        "Airflow se propone para programar el pipeline, ordenar etapas, reintentar fallos y "
        "evitar publicar informes incompletos. No está implementado en el repositorio; la "
        "orquestación actual es el panel ETL y scripts Python.",
    )

    # 10. Conclusión
    h(doc, "9. Conclusión", Pt(14))
    body(
        doc,
        "Cada área tiene la cantidad de objetivos que su operación justifica (sin equilibrar "
        "artificialmente). Varios objetivos generan más de un informe, incluyendo mezclas de "
        "simples y compuestos. Los informes simples salen de la capa operacional y están "
        "implementados en listas y paneles. Los compuestos usan ETL y la base columnar. "
        "Airflow automatizaría ese pipeline cuando se incorpore a producción.",
    )
    return doc


def to_pdf(docx: Path, pdf: Path) -> bool:
    try:
        from docx2pdf import convert

        convert(str(docx), str(pdf))
        return pdf.exists()
    except Exception as e:
        print(f"PDF no generado: {e}")
        return False


def main() -> None:
    doc = build()
    doc.save(DOCX)
    print(f"DOCX: {DOCX}")
    if to_pdf(DOCX, PDF):
        print(f"PDF: {PDF}")
    else:
        print("Exporta el DOCX a PDF desde Word.")


if __name__ == "__main__":
    main()
